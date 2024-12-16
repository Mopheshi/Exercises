import os

from docx import Document

# Create a new Document
doc = Document()

# Title
doc.add_heading('', level=1)

# Sections
doc.add_heading('', level=2)
doc.add_paragraph('')

doc.add_heading('', level=2)
doc.add_paragraph('')

doc.add_heading('', level=2)
doc.add_paragraph('')

doc.add_heading('', level=2)
doc.add_paragraph('')

doc.add_heading('', level=2)
doc.add_paragraph('')

doc.add_heading('', level=2)
doc.add_paragraph('')

doc.add_heading('', level=2)
doc.add_paragraph('')

doc.add_heading('', level=2)
doc.add_paragraph('')

doc.add_paragraph('')

# Ensure the directory exists
file_path = "C:/Users/MOPHE/Desktop/Bilgah App Proposal.docx"
os.makedirs(os.path.dirname(file_path), exist_ok=True)

# Save to a file
doc.save(file_path)
